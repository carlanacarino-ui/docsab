"""
Generador de Documentos A y B
App Streamlit para procesar documentos fuente y generar análisis final.
Usa OpenRouter + Claude Opus / Gemini
"""

import io
import json
import os
import re
import tempfile
import zipfile
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Dict, List

import pandas as pd
import requests
import streamlit as st
from docx import Document
from docx.shared import Pt
from pypdf import PdfReader


# =========================
# Configuración
# =========================

OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")
OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"

MODELS = {
    "Claude Opus 4.7": "anthropic/claude-opus-4.7",
    "Gemini 2.5 Pro": "google/gemini-2.5-pro",
}

DOCUMENT_TYPES = [
    "Ficha de acción formativa",
    "Pliego del lote",
    "Memoria de captación de empresas y perfiles participantes",
    "Memoria de impartición de formación",
    "Memoria de asesoramiento a empresas",
    "Otro / No clasificado",
]


@dataclass
class SourceDoc:
    filename: str
    doc_type: str
    text: str


# =========================
# Lectura de archivos
# =========================

def read_uploaded_file(uploaded_file) -> str:
    """Extrae texto de PDF, DOCX, TXT, XLSX."""
    suffix = Path(uploaded_file.name).suffix.lower()
    raw = uploaded_file.read()

    if suffix == ".pdf":
        reader = PdfReader(io.BytesIO(raw))
        pages = []
        for page in reader.pages:
            pages.append(page.extract_text() or "")
        return "\n\n".join(pages)

    if suffix == ".docx":
        doc = Document(io.BytesIO(raw))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells))
        return "\n".join(parts)

    if suffix in [".txt", ".md"]:
        return raw.decode("utf-8", errors="ignore")

    if suffix in [".xlsx", ".xls"]:
        sheets = pd.read_excel(io.BytesIO(raw), sheet_name=None)
        chunks = []
        for name, df in sheets.items():
            chunks.append(f"## Hoja: {name}\n{df.to_markdown(index=False)}")
        return "\n\n".join(chunks)

    raise ValueError(f"Formato no soportado: {suffix}")


def truncate_text(text: str, limit: int = 45_000) -> str:
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if len(text) <= limit:
        return text
    return text[:limit] + "\n\n[TRUNCADO POR LONGITUD]"


# =========================
# OpenRouter API
# =========================

def call_openrouter(system: str, user: str, model: str, max_tokens: int = 4000) -> str:
    """Llama a OpenRouter con Claude o Gemini."""
    if not OPENROUTER_API_KEY:
        st.error("Falta OPENROUTER_API_KEY. Define la variable de entorno.")
        st.stop()

    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://talentree.es",
    }

    # Combina system y user en un único mensaje
    combined_message = f"{system}\n\n{user}"

    payload = {
        "model": model,
        "messages": [
            {"role": "user", "content": combined_message},
        ],
        "max_tokens": max_tokens,
        "temperature": 0.2,
    }

    response = requests.post(OPENROUTER_URL, json=payload, headers=headers, timeout=120)

    if response.status_code != 200:
        st.error(f"Error OpenRouter: {response.status_code} - {response.text}")
        st.stop()

    data = response.json()

    # Debug: mostrar respuesta si hay error
    if "error" in data:
        st.error(f"Error en respuesta OpenRouter: {data}")
        st.stop()

    if "choices" not in data or not data["choices"]:
        st.error(f"Respuesta OpenRouter inválida (sin 'choices'): {data}")
        st.stop()

    return data["choices"][0]["message"]["content"]


# =========================
# Clasificación
# =========================

def classify_document(filename: str, text: str) -> str:
    lowered = (filename + "\n" + text[:3000]).lower()
    rules = [
        ("ficha", "Ficha de acción formativa"),
        ("acción formativa", "Ficha de acción formativa"),
        ("pliego", "Pliego del lote"),
        ("prescripciones técnicas", "Pliego del lote"),
        ("captación", "Memoria de captación de empresas y perfiles participantes"),
        ("perfiles", "Memoria de captación de empresas y perfiles participantes"),
        ("impartición", "Memoria de impartición de formación"),
        ("formación a los trabajadores", "Memoria de impartición de formación"),
        ("asesoramiento", "Memoria de asesoramiento a empresas"),
    ]
    for needle, label in rules:
        if needle in lowered:
            return label
    return "Otro / No clasificado"


def extract_expediente_and_lote(text: str) -> tuple:
    """Intenta extraer expediente y lote del contenido del documento."""
    expediente = None
    lote = None

    # Busca patrones de expediente: 2021/3120012029/539, etc.
    import re
    exp_match = re.search(r'2021/\d+/\d{3}', text)
    if exp_match:
        expediente = exp_match.group(0)

    # Busca patrón "Lote N" o "Lote N —"
    lote_match = re.search(r'Lote\s+(\d+)', text, re.IGNORECASE)
    if lote_match:
        lote_num = lote_match.group(1)
        # Intenta encontrar el nombre del lote después
        sector_match = re.search(rf'Lote\s+{lote_num}\s*—?\s*([^,\n]+)', text)
        if sector_match:
            lote = f"Lote {lote_num} — {sector_match.group(1).strip()}"
        else:
            lote = f"Lote {lote_num}"

    return expediente, lote


# =========================
# Síntesis
# =========================

def summarize_source(doc: SourceDoc, sector: str, lote: str, expediente: str, model: str) -> str:
    system = """
Eres especialista en evaluación de programas de formación subvencionada, Fundae, sostenibilidad empresarial y elaboración de memorias justificativas.
Tu tarea es convertir documentos fuente heterogéneos en una ficha estructurada, objetiva y reutilizable.
No inventes datos. Si algo no aparece, marca "No consta". Responde siempre en español formal.
""".strip()

    user = f"""
Contexto:
- Sector: {sector}
- Lote: {lote}
- Expediente: {expediente}

Documento: {doc.filename}
Tipo: {doc.doc_type}

Extrae una síntesis estructurada:

1. Identificación del documento
2. Datos cuantitativos relevantes
3. Objetivos
4. Actuaciones realizadas
5. Resultados
6. Indicadores para Documento A
7. Conclusiones para Documento B
8. Carencias o limitaciones

Texto:
{truncate_text(doc.text)}
""".strip()
    return call_openrouter(system, user, model, max_tokens=3000)


def aggregate_summaries(summaries: Dict[str, str], sector: str, lote: str, expediente: str, model: str) -> str:
    system = """
Eres director metodológico de un entregable Fundae.
Integra síntesis parciales en una base maestra coherente para redactar Documento A y Documento B.
No inventes cifras; usa "No consta" cuando proceda.
""".strip()

    joined = "\n\n".join(f"===== {name} =====\n{summary}" for name, summary in summaries.items())

    user = f"""
Contexto:
- Sector: {sector}
- Lote: {lote}
- Expediente: {expediente}

A partir de estas síntesis, crea una BASE MAESTRA estructurada con:

A. Datos generales
B. Inventario de fuentes
C. Requisitos del pliego
D. Caracterización de empresas y participantes
E. Acciones formativas
F. Datos de impartición
G. Datos de captación
H. Datos de asesoramiento
I. Indicadores de eficacia
J. Indicadores de eficiencia
K. Calidad y satisfacción
L. Impacto
M. Fortalezas
N. Debilidades
O. Propuestas preliminares
P. Datos ausentes

Síntesis:
{joined}
""".strip()
    return call_openrouter(system, user, model, max_tokens=6000)


# =========================
# Generación final
# =========================

def generate_document_a(master_base: str, sector: str, lote: str, expediente: str, adjudicataria: str, contratante: str, model: str) -> str:
    system = """
Eres redactor sénior de informes institucionales Fundae.
Redacta un Documento A profesional, analítico y justificativo.
No inventes datos. Incluye tablas en markdown cuando aporten claridad.
""".strip()

    user = f"""
Redacta el DOCUMENTO A completo.

TÍTULO: INFORME DE ANÁLISIS — DOCUMENTO A
Análisis de los resultados, la eficacia y la eficiencia de las actuaciones de formación realizadas

Metadatos:
- Sector: {sector}
- Expediente: {expediente}
- Lote: {lote}
- Entidad adjudicataria: {adjudicataria}
- Entidad contratante: {contratante}

Estructura obligatoria:
1. Resumen ejecutivo
2. Marco de referencia y metodología
3. Caracterización de la población formada
4. Las acciones formativas desarrolladas
5. Análisis de eficacia
6. Análisis de eficiencia
7. Análisis de calidad y satisfacción
8. Análisis de impacto y transferencia
9. Resultados del asesoramiento
10. Síntesis y valoración global

Requisitos:
- Tono formal institucional.
- Incluye tablas de indicadores.
- Distingue datos constatados, estimados y no disponibles.

Base maestra:
{master_base}
""".strip()
    return call_openrouter(system, user, model, max_tokens=8000)


def generate_document_b(master_base: str, doc_a: str, sector: str, lote: str, expediente: str, adjudicataria: str, contratante: str, model: str) -> str:
    system = """
Eres redactor sénior de conclusiones estratégicas, sostenibilidad y mejora continua para Fundae.
Redacta un Documento B complementario, sin repetir el análisis del A.
No inventes datos. Formula propuestas accionables y priorizadas.
""".strip()

    user = f"""
Redacta el DOCUMENTO B completo.

TÍTULO: INFORME DE CONCLUSIONES Y PROPUESTAS — DOCUMENTO B
Conclusiones de la experimentación realizada y propuestas para la mejora y sostenibilidad del proceso

Metadatos:
- Sector: {sector}
- Expediente: {expediente}
- Lote: {lote}
- Entidad adjudicataria: {adjudicataria}
- Entidad contratante: {contratante}

Estructura obligatoria:
1. Introducción
2. Conclusiones de la experimentación
3. Fortalezas del programa
4. Debilidades y limitaciones
5. Propuestas de mejora
6. Sostenibilidad del proceso
7. Hoja de ruta a 36 meses
8. Indicadores para futuras ediciones
9. Cierre y conclusión final

Requisitos:
- Complementa al Documento A, no lo repitas.
- Incluye propuestas quick wins, estratégicas e incrementales.
- Incluye hoja de ruta 0–3, 3–6, 6–12, 12–24 y 24–36 meses.

Base maestra:
{master_base}

Documento A para coherencia:
{truncate_text(doc_a, 30_000)}
""".strip()
    return call_openrouter(system, user, model, max_tokens=8000)


# =========================
# Exportación DOCX
# =========================

def text_to_docx(content: str, title: str, metadata: Dict[str, str]) -> bytes:
    doc = Document()

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)

    doc.add_heading(title, 0)
    for key, value in metadata.items():
        p = doc.add_paragraph()
        p.add_run(f"{key}: ").bold = True
        p.add_run(value or "No consta")

    doc.add_page_break()

    lines = content.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip():
            i += 1
            continue

        # Tabla markdown
        if "|" in line and i + 1 < len(lines) and "|" in lines[i + 1] and "---" in lines[i + 1]:
            table_lines = []
            while i < len(lines) and "|" in lines[i]:
                table_lines.append(lines[i])
                i += 1
            try:
                clean = [l.strip() for l in table_lines if l.strip()]
                if len(clean) >= 2:
                    rows = []
                    for line in clean:
                        if re.match(r"^\|?\s*:?-{3,}:?", line):
                            continue
                        cells = [c.strip() for c in line.strip("|").split("|")]
                        rows.append(cells)
                    if rows:
                        cols = max(len(r) for r in rows)
                        table = doc.add_table(rows=len(rows), cols=cols)
                        table.style = "Table Grid"
                        for ri, row in enumerate(rows):
                            for ci in range(cols):
                                table.cell(ri, ci).text = row[ci] if ci < len(row) else ""
            except:
                pass
            continue

        # Encabezados markdown
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif re.match(r"^\d+\.\s+", line):
            doc.add_heading(line.strip(), level=1)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)
        i += 1

    doc.add_paragraph("◆ ◆ ◆")
    doc.add_paragraph("Fin del documento")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def make_zip(files: Dict[str, bytes]) -> bytes:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for filename, data in files.items():
            zf.writestr(filename, data)
    return buffer.getvalue()


# =========================
# Streamlit UI
# =========================

st.set_page_config(page_title="Generador Documentos A/B", layout="wide")
st.title("📄 Generador de Documento A y Documento B")
st.caption("Pipeline: extracción → síntesis estructurada → generación final .docx")

with st.sidebar:
    st.header("⚙️ Configuración")

    model_name = st.selectbox("Modelo LLM", list(MODELS.keys()))
    model = MODELS[model_name]

    st.divider()
    st.subheader("📋 Datos del expediente")
    sector = st.text_input("Sector", value="METAL")
    expediente = st.text_input("Expediente", value="2021/3120012029/539")
    lote = st.text_input("Lote", value="Lote 6 — Metal")
    adjudicataria = st.text_input("Entidad adjudicataria", value="MainJobs Internacional")
    contratante = st.text_input("Entidad contratante", value="Fundae")

uploaded_files = st.file_uploader(
    "📤 Sube los documentos fuente (PDF, DOCX, XLSX, TXT)",
    type=["pdf", "docx", "txt", "xlsx", "xls"],
    accept_multiple_files=True,
)

if uploaded_files:
    st.subheader("Paso 1️⃣ — Extracción y clasificación")

    docs: List[SourceDoc] = []
    rows = []
    for f in uploaded_files:
        try:
            text = read_uploaded_file(f)
            detected = classify_document(f.name, text)
            docs.append(SourceDoc(filename=f.name, doc_type=detected, text=text))
            rows.append({"Archivo": f.name, "Tipo detectado": detected, "Caracteres": len(text)})
        except Exception as exc:
            st.error(f"Error leyendo {f.name}: {exc}")

    if rows:
        df = pd.DataFrame(rows)
        st.dataframe(df, use_container_width=True, hide_index=True)

        # Verificar coherencia de expediente y lote
        detected_expedientes = set()
        detected_lotes = set()
        for d in docs:
            exp, lote = extract_expediente_and_lote(d.text)
            if exp:
                detected_expedientes.add(exp)
            if lote:
                detected_lotes.add(lote)

        if detected_expedientes and expediente not in detected_expedientes:
            st.warning(f"⚠️ Los documentos contienen expediente(s): {', '.join(detected_expedientes)}\nPero has configurado: {expediente}")

        if detected_lotes:
            st.info(f"📋 Lotes detectados en documentos: {', '.join(detected_lotes)}")

        st.info("📝 Puedes corregir manualmente el tipo de cada documento antes de generar.")

        corrected_docs: List[SourceDoc] = []
        cols = st.columns(2)
        for idx, d in enumerate(docs):
            col_idx = idx % 2
            with cols[col_idx]:
                st.write(f"**{d.filename}**")
                selected = st.selectbox(
                    "Tipo",
                    DOCUMENT_TYPES,
                    index=DOCUMENT_TYPES.index(d.doc_type) if d.doc_type in DOCUMENT_TYPES else len(DOCUMENT_TYPES) - 1,
                    key=f"doctype_{idx}",
                    label_visibility="collapsed",
                )
            corrected_docs.append(SourceDoc(d.filename, selected, d.text))

        run = st.button("▶️ Ejecutar los 3 pasos", type="primary", use_container_width=True)

        if run:
            progress = st.progress(0)
            status = st.empty()

            # Paso 2: Síntesis
            st.subheader("Paso 2️⃣ — Síntesis estructurada")
            summaries: Dict[str, str] = {}

            for idx, d in enumerate(corrected_docs):
                status.text(f"⏳ Sintetizando: {d.filename}...")
                summaries[d.filename] = summarize_source(d, sector, lote, expediente, model)
                progress.progress(int((idx + 1) / len(corrected_docs) * 30))

            with st.expander("📋 Ver síntesis parciales"):
                for name, summary in summaries.items():
                    st.markdown(f"**{name}**")
                    st.markdown(summary[:500] + "...")

            # Base maestra
            status.text("🔗 Construyendo base maestra...")
            master_base = aggregate_summaries(summaries, sector, lote, expediente, model)
            progress.progress(50)

            with st.expander("📊 Ver base maestra"):
                st.markdown(master_base[:1000] + "...")

            # Paso 3: Documentos finales
            st.subheader("Paso 3️⃣ — Generación final")

            status.text("✍️ Redactando Documento A...")
            doc_a_text = generate_document_a(master_base, sector, lote, expediente, adjudicataria, contratante, model)
            progress.progress(75)

            status.text("✍️ Redactando Documento B...")
            doc_b_text = generate_document_b(master_base, doc_a_text, sector, lote, expediente, adjudicataria, contratante, model)
            progress.progress(90)

            metadata = {
                "Sector": sector,
                "Expediente": expediente,
                "Lote": lote,
                "Entidad adjudicataria": adjudicataria,
                "Entidad contratante": contratante,
                "Fecha": datetime.now().strftime("%d/%m/%Y %H:%M"),
                "Modelo": model_name,
            }

            doc_a_bytes = text_to_docx(doc_a_text, "INFORME DE ANÁLISIS — DOCUMENTO A", metadata)
            doc_b_bytes = text_to_docx(doc_b_text, "INFORME DE CONCLUSIONES Y PROPUESTAS — DOCUMENTO B", metadata)

            zip_bytes = make_zip({
                "Documento_A_Analisis_resultados_eficacia_eficiencia.docx": doc_a_bytes,
                "Documento_B_Conclusiones_propuestas_mejora_sostenibilidad.docx": doc_b_bytes,
                "base_maestra.txt": master_base.encode("utf-8"),
            })

            progress.progress(100)
            status.success("✅ ¡Documentos generados!")

            st.divider()
            col_a, col_b, col_zip = st.columns(3)
            with col_a:
                st.download_button(
                    "📥 Documento A",
                    data=doc_a_bytes,
                    file_name="Documento_A_Analisis_resultados_eficacia_eficiencia.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            with col_b:
                st.download_button(
                    "📥 Documento B",
                    data=doc_b_bytes,
                    file_name="Documento_B_Conclusiones_propuestas_mejora_sostenibilidad.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            with col_zip:
                st.download_button(
                    "📦 TODO en ZIP",
                    data=zip_bytes,
                    file_name="documentos_generados.zip",
                    mime="application/zip",
                    use_container_width=True,
                )

            st.divider()
            st.subheader("👀 Vista previa")
            tab_a, tab_b = st.tabs(["Documento A", "Documento B"])
            with tab_a:
                st.markdown(doc_a_text)
            with tab_b:
                st.markdown(doc_b_text)

else:
    st.info("👆 Sube entre 1 y 11 documentos para empezar")
